"""Audit OE04 Word + Excel vs ficha de lançamento."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

# Resolve M5 OE4 folder without embedding emoji in source
PEF = Path.home() / "OneDrive - Universidade de Aveiro" / "PEF"
OE4_DIR = next(
    p for p in (PEF / "M5✔️").iterdir() if p.is_dir() and "OE4" in p.name
)
XLSX = OE4_DIR / "OE04_G18_260525.xlsx"
DOCX = OE4_DIR / "OE04_G18_260525.docx"

EXPECTED_SHEETS = [
    "00_PRESSUPOSTOS",
    "01_Pre_Projeto",
    "02_Mapa_Investimento",
    "03_Estrutura_Capital",
    "04_Plano_Financ",
    "05_SD_BancoHub",
    "06_SD_BEI",
    "07_PT2030",
    "08_SD_Existente",
    "09_Pos_Projeto",
]

# Key figures from motor / prompt (tolerance for rounding)
KEY_FIGURES = {
    "capex_total": 6_000_000,
    "banco_hub": 3_000_000,
    "bei": 1_500_000,
    "pt2030": 2_700_000,
    "aft_2029": 3_550_110,
    "autonomia_min_pct": 0.30,
}


def num(s) -> float | None:
    if s is None:
        return None
    if isinstance(s, (int, float)):
        return float(s)
    t = str(s).strip().replace("\u00a0", " ").replace("€", "").replace("%", "")
    t = t.replace(".", "").replace(",", ".") if re.search(r"\d,\d", t) else t.replace(",", "")
    m = re.search(r"-?[\d.]+", t.replace(" ", ""))
    return float(m.group()) if m else None


def audit_xlsx(path: Path) -> dict:
    wb = load_workbook(path, data_only=True)
    out: dict = {"path": str(path), "sheets": wb.sheetnames, "issues": [], "checks": {}}

    missing = [s for s in EXPECTED_SHEETS if s not in wb.sheetnames]
    if missing:
        out["issues"].append(f"Folhas em falta: {missing}")

    if "00_PRESSUPOSTOS" in wb.sheetnames:
        ws = wb["00_PRESSUPOSTOS"]
        labels = {}
        for row in ws.iter_rows(min_row=1, max_row=80, min_col=1, max_col=2):
            a, b = row[0].value, row[1].value
            if a and b is not None:
                labels[str(a).strip().lower()] = b
        for key, substr, field in [
            ("capex", "capex base total", "capex_total"),
            ("banco", "banco hub — capital", "banco_hub"),
            ("bei", "bei — capital", "bei"),
            ("pt", "pt2030 — montante", "pt2030"),
            ("aft", "aft hub líquido", "aft_2029"),
        ]:
            val = next((v for k, v in labels.items() if substr in k), None)
            if val is not None:
                exp = KEY_FIGURES[field]
                ok = abs(float(val) - exp) < 1
                out["checks"][field] = {"excel": val, "expected": exp, "ok": ok}
                if not ok:
                    out["issues"].append(f"00_PRESSUPOSTOS {field}: {val} != {exp}")

    if "09_Pos_Projeto" in wb.sheetnames:
        ws = wb["09_Pos_Projeto"]
        aut_vals = []
        for row in ws.iter_rows():
            for c in row:
                if c.value is None:
                    continue
                lab = str(c.value).lower()
                if "autonomia" in lab:
                    # look right for %
                    for ncell in row[1:8]:
                        v = num(ncell.value)
                        if v is not None and 0 < v <= 1.5:
                            aut_vals.append(v)
        if aut_vals:
            mn = min(aut_vals)
            out["checks"]["autonomia_min"] = {
                "min_found": mn,
                "ok": mn >= KEY_FIGURES["autonomia_min_pct"] - 0.001,
            }
            if mn < KEY_FIGURES["autonomia_min_pct"] - 0.001:
                out["issues"].append(f"Autonomia < 30%: min={mn:.1%}")

    # Debt service sheets non-empty
    for sn in ["05_SD_BancoHub", "06_SD_BEI", "07_PT2030", "08_SD_Existente"]:
        if sn in wb.sheetnames:
            ws = wb[sn]
            filled = sum(1 for r in ws.iter_rows() for c in r if c.value not in (None, ""))
            out["checks"][f"rows_{sn}"] = filled
            if filled < 5:
                out["issues"].append(f"{sn} parece vazio ({filled} células)")

    wb.close()
    return out


def all_doc_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def audit_docx(path: Path) -> dict:
    doc = Document(path)
    text = all_doc_text(doc)
    out: dict = {"path": str(path), "issues": [], "checks": {}, "sections": {}}

    # Filename convention
    if path.name != "OE04_G18_260525.docx":
        out["issues"].append(f"Nome ficheiro: {path.name}")

    # Format hints (Arial 10 — sample first runs)
    fonts = set()
    for p in doc.paragraphs[:50]:
        for r in p.runs:
            if r.font.name:
                fonts.add(r.font.name)
    out["checks"]["fonts_sample"] = sorted(fonts)

    # Required thematic blocks (OE4 ficha)
    required_topics = [
        ("pre_projeto", r"pr[eé].?projeto|situa[cç][aã]o.*equil[ií]brio.*pr[eé]"),
        ("mapa_investimento", r"mapa de investimento|investimento.*projeto|capex"),
        ("estrutura_capital", r"estrutura de capital|crit[eé]rios"),
        ("plano_financiamento", r"plano de financiamento|fontes de financiamento"),
        ("duas_fontes_alheio", r"(leasing|empr[eé]stimo|subs[ií]dio|bei|banco hub|pt\s*2030).*(leasing|empr[eé]stimo|subs[ií]dio|bei|banco|pt\s*2030)", re.I | re.S),
        ("servico_divida", r"servi[cç]o da d[ií]vida|mapas? de servi[cç]o"),
        ("pos_projeto", r"p[oó]s.?projeto|autonomia financeira"),
        ("fiscal", r"fiscal|parafiscal|economias fiscais|juros dedut"),
        ("legal_pt2030", r"portugal\s*2030|elegibilidade|enquadramento jur[ií]dico"),
        ("ferramenta_digital", r"ferramenta digital|excel|anexo"),
        ("pressupostos", r"pressupostos"),
    ]
    tl = text.lower()
    for name, pat, *flags in required_topics:
        fl = flags[0] if flags else 0
        found = bool(re.search(pat, text, fl))
        out["sections"][name] = found
        if not found and name not in ("duas_fontes_alheio",):  # checked separately
            if name == "duas_fontes_alheio":
                pass
            else:
                out["issues"].append(f"Secção possivelmente em falta: {name}")

    # Two distinct debt sources — explicit mentions
    sources = []
    for kw in ["banco hub", "bei", "pt2030", "pt 2030", "leasing", "empréstimo", "empréstimo bancário"]:
        if kw in tl:
            sources.append(kw)
    out["checks"]["financing_sources_mentioned"] = list(dict.fromkeys(sources))
    if len(out["checks"]["financing_sources_mentioned"]) < 2:
        out["issues"].append("Menos de 2 fontes de capital alheio identificadas no texto")

    # Autonomia 30%
    aut_match = re.findall(r"autonomia[^.\n]{0,80}?([\d]+[,.][\d]+)\s*%", text, re.I)
    aut_ok = any(num(x) and num(x) >= 29.5 for x in aut_match)
    out["checks"]["autonomia_30_mentioned"] = aut_ok
    if not aut_ok:
        out["issues"].append("Não encontrei menção clara a autonomia ≥30%")

    # Excel reference coherence
    if "OE04_G18_260525.xlsx" in text:
        out["checks"]["excel_filename_ref"] = "OE04_G18_260525.xlsx"
    elif "FINAL_REV" in text:
        out["issues"].append("Word referencia FINAL_REV em vez de OE04_G18_260525.xlsx")
    else:
        out["issues"].append("Word não referencia nome Excel submissão")

    # Key numbers in Word
    for label, val in [
        ("6.000.000", 6_000_000),
        ("3.000.000", 3_000_000),
        ("2.700.000", 2_700_000),
        ("3.550.110", 3_550_110),
    ]:
        variants = [label, label.replace(".", " ")]
        out["checks"][f"mentions_{label}"] = any(v in text for v in variants)

    # Page count estimate (body only rough)
    out["checks"]["paragraph_count"] = len(doc.paragraphs)

    return out


def cross_check(xlsx_audit: dict, doc_audit: dict) -> list[str]:
    issues = []
    xc = xlsx_audit.get("checks", {})
    dc = doc_audit.get("checks", {})
    if xc.get("capex_total", {}).get("excel") and not dc.get("mentions_6.000.000"):
        issues.append("CAPEX 6M no Excel mas não óbvio no Word")
    if xc.get("aft_2029", {}).get("excel") and not dc.get("mentions_3.550.110"):
        issues.append("AFT 3.550.110 no Excel mas não no Word")
    if dc.get("excel_filename_ref") and "OE04_G18_260525.xlsx" not in str(xlsx_audit.get("path", "")):
        issues.append("Nome Excel Word vs ficheiro analisado")
    return issues


def main() -> int:
    print("=== OE04 AUDIT ===")
    print("Excel:", XLSX)
    print("Word:", DOCX)
    if not XLSX.exists() or not DOCX.exists():
        print("FICHEIROS EM FALTA")
        return 1

    xa = audit_xlsx(XLSX)
    da = audit_docx(DOCX)
    cross = cross_check(xa, da)

    print("\n--- EXCEL ---")
    print("Folhas:", xa["sheets"])
    print("Checks:", xa["checks"])
    for i in xa["issues"]:
        print("ISSUE:", i)

    print("\n--- WORD ---")
    print("Secções:", da["sections"])
    print("Checks:", {k: v for k, v in da["checks"].items() if k != "fonts_sample"})
    print("Fontes (amostra):", da["checks"].get("fonts_sample"))
    for i in da["issues"]:
        print("ISSUE:", i)

    print("\n--- COERÊNCIA EXCEL/WORD ---")
    for i in cross:
        print("ISSUE:", i)
    if not cross:
        print("Sem inconsistências críticas de números/nomes detectadas pelo script.")

    print("\n--- FICHA OE04 (checklist manual automatizada) ---")
    checklist = [
        ("Análise equilíbrio pré-projeto", da["sections"].get("pre_projeto")),
        ("Mapa investimento", da["sections"].get("mapa_investimento")),
        ("Critérios estrutura capital", da["sections"].get("estrutura_capital")),
        ("Plano financiamento + ≥2 fontes alheio", da["sections"].get("plano_financiamento") and len(da["checks"].get("financing_sources_mentioned", [])) >= 2),
        ("Mapas serviço dívida", da["sections"].get("servico_divida") and xa["checks"].get("rows_05_SD_BancoHub", 0) > 5),
        ("Equilíbrio pós-projeto + 30% autonomia", da["sections"].get("pos_projeto") and xa["checks"].get("autonomia_min", {}).get("ok", False)),
        ("Fiscais/parafiscais", da["sections"].get("fiscal")),
        ("Legais PT2030", da["sections"].get("legal_pt2030")),
        ("Ferramenta digital (Excel anexo)", da["sections"].get("ferramenta_digital")),
        ("10 folhas Excel estruturadas", len(xa["issues"]) == 0 or "Folhas em falta" not in str(xa["issues"])),
    ]
    for label, ok in checklist:
        print(f"  [{'OK' if ok else '??'}] {label}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
