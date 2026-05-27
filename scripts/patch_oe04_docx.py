"""Apply OE4 Word corrections to OE04_G18_260525.docx."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOC_PATH = Path(__file__).resolve().parents[1] / "OE04_G18_260525.docx"


def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_in_paragraph(paragraph: Paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    new_text = paragraph.text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)
    return True


def replace_everywhere(doc: Document, replacements: list[tuple[str, str]]) -> int:
    n = 0
    for old, new in replacements:
        for p in iter_all_paragraphs(doc):
            if old in p.text and replace_in_paragraph(p, old, new):
                n += 1
    return n


def set_cell_text(cell, value: str) -> None:
    if cell.paragraphs:
        cell.paragraphs[0].text = value
        for p in cell.paragraphs[1:]:
            p.text = ""
    else:
        cell.text = value


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p_el = OxmlElement("w:p")
    paragraph._p.addnext(new_p_el)
    new_para = Paragraph(new_p_el, paragraph._parent)
    new_para.text = text
    return new_para


def patch_tables(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "395.250" in cell.text:
                    set_cell_text(cell, cell.text.replace("395.250", "305.250"))
                if "3.595.145" in cell.text:
                    set_cell_text(cell, cell.text.replace("3.595.145", "3.550.110"))

        header = " ".join(c.text for c in table.rows[0].cells)
        if "Cash-in" in header and "Accrual" in header:
            cash_col = saldo_col = None
            for ci, c in enumerate(table.rows[0].cells):
                h = c.text
                if "Cash-in" in h:
                    cash_col = ci
                if "Saldo" in h and "Diferido" in h:
                    saldo_col = ci
            for row in table.rows[1:]:
                if row.cells[0].text.strip() == "2027":
                    if cash_col is not None and row.cells[cash_col].text.strip() in ("0", ""):
                        set_cell_text(row.cells[cash_col], "2.700.000")
                    if saldo_col is not None and "2.382" in row.cells[saldo_col].text:
                        set_cell_text(row.cells[saldo_col], "2.065.468")

        if table.rows and len(table.rows[0].cells) >= 5:
            hrow = " ".join(c.text for c in table.rows[0].cells)
            if "Amort" in hrow and "Saldo" in hrow:
                for row in table.rows[1:]:
                    if row.cells[0].text.strip() == "2027" and "3.000" in row.cells[1].text:
                        if row.cells[4].text.strip() in ("—", "-", ""):
                            set_cell_text(row.cells[4], "2.700")

    rfai_data = {
        "2025": ("600.000", "0", "0", "0"),
        "2026": ("600.000", "88.568", "44.284", "44.284"),
        "2027": ("555.716", "124.044", "62.022", "62.022"),
        "2028": ("493.694", "148.579", "74.290", "74.290"),
        "2029": ("419.405", "163.769", "81.884", "81.884"),
        "2030": ("337.520", "192.732", "96.366", "96.366"),
        "2031": ("241.154", "228.064", "114.032", "114.032"),
        "2032": ("127.122", "237.438", "118.719", "118.719"),
        "2033": ("8.403", "247.140", "8.403", "8.403"),
        "2034": ("0", "279.162", "0", "0"),
    }
    for table in doc.tables:
        hdr = " ".join(c.text for row in table.rows[:2] for c in row.cells)
        if "RFAI Aplicado" in hdr or ("Crédito Remanescente" in hdr and "IRC Bruto" in hdr):
            for row in table.rows[1:]:
                yr = row.cells[0].text.strip()
                if yr in rfai_data:
                    for ci, v in enumerate(rfai_data[yr]):
                        if ci + 1 < len(row.cells):
                            set_cell_text(row.cells[ci + 1], v)


def add_notes(doc: Document) -> None:
    pt2030_cash_note = (
        "O cash-in de 2.700.000 € em 2027 é aplicado à amortização antecipada do empréstimo "
        "Banco Hub (Folha 05_MSD_BancoHub), não duplicando o financiamento do CAPEX."
    )
    nd_note = (
        "Os rácios de AF em 2030–2034 no Excel são extrapolações simplificadas (n.d.) e não "
        "devem ser citados como prova adicional do cumprimento da OE4."
    )
    vala_note = (
        "Nota sobre VALA e horizonte temporal. A OE4 exige o plano de financiamento e AF ≥ 30%, "
        "com projeções patrimoniais robustas em 2025–2029. O horizonte 2030–2034 no Excel suporta "
        "FCF, VAL, TIR, payback, RFAI e valor terminal, com extrapolação simplificada. A síntese VALA "
        "(Folha 11_VALA) é indicativa e não se reproduz neste relatório por não ser requisito da OE4; "
        "o rigor patrimonial completo até 2034 será desenvolvido na iteração M6."
    )
    vala_anexo = (
        "Nota VALA: não apresentada no relatório da OE4; ver Folha 11_VALA e limitações em 2030–2034."
    )

    anchors = [
        ("perfaz os 2.700", pt2030_cash_note),
        ("perfaz os 2.700.000", pt2030_cash_note),
        ("Tabela 6a — Âmbito Metodológico", nd_note),
        ("Esta limitação não altera a conclusão de viabilidade da OE4", vala_note),
        ("O VAL foi calculado pela soma dos fluxos", vala_anexo),
    ]
    for anchor, note in anchors:
        for p in doc.paragraphs:
            if anchor in p.text:
                nxt = p._element.getnext()
                if nxt is not None:
                    from docx.oxml.text.paragraph import CT_P

                    if isinstance(nxt, CT_P):
                        para = Paragraph(nxt, p._parent)
                        if note[:35] in para.text:
                            break
                insert_paragraph_after(p, note)
                break


def main() -> int:
    path = Path(sys.argv[1]) if len(sys.argv) > 1 else DOC_PATH
    doc = Document(str(path))

    reps = [
        ("OE04_G18_260525_FINAL_REV.xlsx", "OE04_G18_260525.xlsx"),
        ("WACC específico do projeto em 7,30% nominal", "WACC específico do projeto em 6,3% nominal"),
        ("WACC resultante é de 7,30% nominal", "WACC adoptado para o VAL e a TIR é de 6,3% nominal"),
        ("descontados ao WACC de 7,30%", "descontados ao WACC de 6,3%"),
        ("WACC de 7,30%", "WACC de 6,3%"),
        ("Crédito 10% × 6 M€ = 600 m€ deduzido", "Crédito 10% × 6 M€ = 600.000 € deduzido"),
        ("600 m€ deduzido", "600.000 € deduzido"),
        ("Dívida Líq./EBITDA ≤ 4,0×", "Dívida Líq./EBITDA ≤ 3,5×"),
        ("317 m€ em 2026, 2027 e 2028", "317.266 € em 2026, 2027 e 2028"),
        ("decrescendo para 288 m€ em 2029", "decrescendo para 288.016 € em 2029"),
        ("6.000 m€ em todos os anos", "6.000.000 € em todos os anos"),
        ("saldo diferido remanescente de 782 m€ em 2034", "saldo diferido remanescente de 782.012 € em 2034"),
        ("totalizam 1.918 m€)", "totalizam 1.918.000 €)"),
        ("450 m Eur (2034)", "450.000 € (2034)"),
        (
            "o crédito RFAI totaliza 110.604 € absorvidos entre 2026 e 2029 e 600.000 € no horizonte completo de análise até 2034",
            "o crédito RFAI totaliza aproximadamente 262.000 € absorvidos entre 2026 e 2029 e 600.000 € no horizonte completo de análise até 2034",
        ),
        (
            "110.604 € absorvidos em 2026–2029; 600.000 € absorvidos no horizonte 2026–2034",
            "aprox. 262.000 € absorvidos em 2026–2029; 600.000 € absorvidos no horizonte 2026–2034",
        ),
        (
            "A dívida residual de 1.350.000 € (BH + BEI) não é deduzida do terminal value.",
            "No FCFF não se deduz dívida ao valor terminal; em 2034 remanescente apenas a Linha BEI (450.000 €), com o Banco Hub liquidado em 2028.",
        ),
        (
            "o projeto apresentar VAL positivo mesmo num cenário alternativo com exclusão total dos fundos públicos",
            "a viabilidade deteriorar-se materialmente em cenário sem apoio público, reforçando o papel do PT2030 na desalavancagem (não na duplicação de fontes)",
        ),
        (
            "tomando como referência oficial os outputs do motor analítico validados na Folha 00: VAL positivo, TIR superior ao WACC e payback compatível com o perfil operacional do Hub. A Folha 10 constitui uma verificação cruzada em Excel com metodologia simplificada, cuja harmonização integral com o motor será desenvolvida em M6.",
            "com base na Folha 10_FCF_Viabilidade do Excel anexo: VAL ≈ 5,08 M€, TIR ≈ 31,3%, payback simples ≈ 3,18 anos, descontados ao WACC de 6,3%. A Folha 00_Pressupostos consolida parâmetros; o motor Python serviu apenas de validação cruzada. A harmonização integral Excel–Python e a extensão patrimonial até 2034 ficam para M6.",
        ),
        (
            "servindo apenas de suporte à auditoria dos resultados obtidos no modelo digital e no presente relatório técnico.",
            "servindo apenas de suporte à auditoria dos resultados obtidos no modelo digital e no presente relatório técnico. A viabilidade económica (VAL, TIR e payback) reporta-se à Folha 10_FCF_Viabilidade (WACC 6,3%). A decomposição VALA (Folha 11_VALA) é complemento M6 e não integra o âmbito formal desta OE4.",
        ),
        (
            "A ferramenta formalmente submetida em anexo é o ficheiro Excel com a designação OE04_G18_260525_FINAL_REV.xlsx",
            "A ferramenta formalmente submetida em anexo é o ficheiro Excel com a designação OE04_G18_260525.xlsx",
        ),
        (
            "Ficheiro Excel anexo: OE04_G18_260525_FINAL_REV.xlsx",
            "Ficheiro Excel anexo: OE04_G18_260525.xlsx",
        ),
    ]
    n = replace_everywhere(doc, reps)
    patch_tables(doc)
    add_notes(doc)
    doc.save(str(path))
    print(f"Replacements in paragraphs/cells: {n}+")
    print(f"Saved: {path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
